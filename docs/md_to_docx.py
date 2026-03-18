#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将本合同 Markdown 转为 Word docx。"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置单元格边框。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("start", "top", "end", "bottom", "left", "right"):
        edge_attr = kwargs.get(edge)
        if edge_attr:
            tag = "w:{}".format(edge)
            element = OxmlElement(tag)
            element.set(qn("w:val"), edge_attr.get("val", "single"))
            element.set(qn("w:sz"), str(edge_attr.get("sz", 4)))
            element.set(qn("w:color"), edge_attr.get("color", "000000"))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def add_paragraph(doc, line):
    line = line.strip()
    if not line:
        return
    # **xxx** -> 保留为普通段落（docx 里可再手动加粗）
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    # 若整行是加粗标签（如 **甲方**：），则加粗到冒号
    if line.startswith("**") and "**" in line[2:]:
        m = re.match(r'\*\*(.+?)\*\*(\s*[：:].*)', line)
        if m:
            run_b = p.add_run(m.group(1) + m.group(2))
            run_b.bold = True
            run_b.font.size = Pt(10.5)
            run_b.font.name = "宋体"
            run_b._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            return
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

def main():
    base = Path(__file__).resolve().parent
    md_path = base / "软件开发合同-道路设计OA系统.md"
    out_path = base / "软件开发合同-道路设计OA系统.docx"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "# 软件开发合同":
            p = doc.add_heading("软件开发合同", level=0)
            p.alignment = 1  # CENTER
            i += 1
            continue

        if stripped.startswith("## "):
            title = stripped[3:].strip()
            doc.add_paragraph()
            p = doc.add_heading(title, level=1)
            for run in p.runs:
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                run.font.size = Pt(12)
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # Markdown 表格
        if stripped.startswith("|") and "|" in stripped[1:]:
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().split("|")[1:-1]]
                if row:
                    rows.append(row)
                i += 1
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        if ci < len(table.rows[ri].cells):
                            table.rows[ri].cells[ci].text = cell_text
                            for p in table.rows[ri].cells[ci].paragraphs:
                                for run in p.runs:
                                    run.font.size = Pt(10.5)
                                    run.font.name = "宋体"
                                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                doc.add_paragraph()
            continue

        add_paragraph(doc, line)
        i += 1

    doc.save(out_path)
    print("已生成:", out_path)

if __name__ == "__main__":
    main()
